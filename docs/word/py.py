'''
** []() $$ не працює точно
зжирає розділи чогось
'''


import re
from docx import Document
from docx.shared import Inches
from datetime import datetime

def parse_metadata(md_content):
    """
    Парсить метадані з Markdown файлу (між --- і ---).
    """
    metadata = {}
    metadata_section = re.search(r"---\n(.*?)\n---", md_content, re.DOTALL)
    if metadata_section:
        lines = metadata_section.group(1).splitlines()
        for line in lines:
            key, value = map(str.strip, line.split(":", 1))
            metadata[key.lower()] = value.strip('"')
    return metadata

def apply_style(doc, text, style_name):
    """
    Додає текст до документа із заданим стилем.
    """
    paragraph = doc.add_paragraph(text)
    paragraph.style = style_name

def replace_placeholders(doc, metadata):
    """
    Замінює маркери (наприклад, {AUTHOR}, {DATE}) у шаблонному документі на значення з метаданих.
    """
    for paragraph in doc.paragraphs:
        for key, value in metadata.items():
            placeholder = f"{{{key.upper()}}}"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)

def process_markdown_file(md_file_path, doc, metadata=None):
    """
    Обробляє один Markdown файл і додає його контент у документ.
    """
    # Читаємо вміст Markdown файлу
    with open(md_file_path, 'r', encoding='utf-8') as md_file:
        md_content = md_file.read()

    # Якщо є метадані в Markdown файлі, замінюємо маркери
    if metadata:
        replace_placeholders(doc, metadata)

    # Очищаємо Markdown текст від метаданих
    content_without_metadata = re.sub(r"---\n.*?\n---", "", md_content, flags=re.DOTALL)

    # Додаємо текст до документа
    lines = content_without_metadata.splitlines()
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue  # Пропускаємо порожні рядки

        if line.startswith("# "):  # Заголовок H1
            apply_style(doc, line[2:], "Heading 1")
        elif line.startswith("## "):  # Заголовок H2
            apply_style(doc, line[3:], "Heading 2")
        elif line.startswith("### "):  # Заголовок H3
            apply_style(doc, line[4:], "Heading 3")
        elif re.match(r"\d+\. ", line):  # Нумерований список
            apply_style(doc, line[line.find(".") + 2:], "List Number")
        elif line.startswith("* "):  # Маркований список
            apply_style(doc, line[2:], "List Bullet")
        elif re.match(r"!\[.*?\]\((.*?)\)", line):  # Картинки
            match = re.search(r"!\[.*?\]\((.*?)\)", line)
            if match:
                image_path = match.group(1)
                try:
                    doc.add_picture(image_path, width=Inches(5))
                except Exception as e:
                    apply_style(doc, f"[Помилка додавання зображення: {image_path}]", "Normal")
        elif re.match(r"\$.*?\$", line):  # Інлайн формули
            formula = re.sub(r"\$(.*?)\$", r"[Formula: \1]", line)
            apply_style(doc, formula, "Normal")
        elif re.match(r"\$\$.*?\$\$", line, re.DOTALL):  # Блок формули
            formula = re.sub(r"\$\$(.*?)\$\$", r"[Block Formula: \1]", line, flags=re.DOTALL)
            apply_style(doc, formula, "Normal")
        elif re.match(r"```", line):  # Блок коду
            code_block = []
            i += 1
            while i < len(lines) and not re.match(r"```", lines[i]):
                code_block.append(lines[i])
                i += 1
            apply_style(doc, "\n".join(code_block), "Code")
        elif re.match(r"`.*?`", line):  # Інлайн код
            inline_code = re.sub(r"`(.*?)`", r"[Code: \1]", line)
            apply_style(doc, inline_code, "Normal")
        elif re.match(r"\[.*?\]\(.*?\)", line):  # Посилання
            link = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1 (\2)", line)
            apply_style(doc, link, "Normal")
        elif re.match(r"\*\*.*?\*\*", line):  # Жирний текст
            bold_text = re.sub(r"\*\*(.*?)\*\*", r"[Bold: \1]", line)
            apply_style(doc, bold_text, "Normal")
        else:  # Звичайний текст
            apply_style(doc, line, "Normal")

def convert_md_structure_to_docx_with_styles(template_docx_path, output_docx_path):
    """
    Конвертує список Markdown файлів із README.md у один .docx.
    """
    # Читаємо вміст README.md
    with open("README.md", 'r', encoding='utf-8') as structure_file:
        structure_content = structure_file.read()

    # Знаходимо посилання на файли (формат: [Назва](шлях))
    md_files = re.findall(r"\[.*?\]\((.*?)\)", structure_content)
    if not md_files:
        print("README.md не містить файлів для обробки.")
        return

    # Завантажуємо шаблонний документ
    doc = Document(template_docx_path)

    metadata = parse_metadata(structure_content)

    # Замінюємо маркери у шаблоні (якщо є загальні метадані)
    replace_placeholders(doc, metadata)

    # Обробляємо кожен Markdown файл
    for md_file in md_files:
        process_markdown_file(md_file, doc, metadata)
        doc.add_paragraph()  # Додаємо порожній абзац між файлами

    # Зберігаємо новий документ
    doc.save(output_docx_path)
    print(f"Файл {output_docx_path} створено успішно.")

# Використання
template_docx_path = "ref/ref.docx"  # Вкажіть шлях до шаблонного .docx
output_docx_path = f"docx/{datetime.now().strftime("%m.%d.%H.%M")}_out.docx"

convert_md_structure_to_docx_with_styles(template_docx_path, output_docx_path)
